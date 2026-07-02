"""import os, sys
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from resume_parser.jd_parser import JDParser
import numpy as np
from ranker.ranker import Ranker
from pprint import pprint
import json
from resume_parser.candidate_parser import CandidateParser
from normalizer.candidate_normalizer import CandidateNormalizer
from matcher.skill_matcher import SkillMatcher
from matcher.experience_matcher import ExperienceMatcher
from matcher.embedding_matcher import EmbeddingMatcher
from scorer.scorer import Scorer


class NumpyEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.ndarray):
            return obj.tolist()  # Converts the numpy array to a standard Python list
        if isinstance(obj, (np.int64, np.int32)):
            return int(obj)      # Handles numpy integers too, just in case
        if isinstance(obj, (np.float64, np.float32)):
            return float(obj)    # Handles numpy floats
        return super().default(obj)


doc = Document("D:\Resume Parser and Ranker - Hackathon\Project\data\job_description.docx")

text = "\n".join(

    paragraph.text

    for paragraph in doc.paragraphs
)

jd_parser = JDParser()

jd = jd_parser.parse(text)

pprint(jd)


# Read original candidate dataset
candidates = []
with open("D:\Resume Parser and Ranker - Hackathon\Project\data\candidates.jsonl", "r", encoding="utf-8") as f:
    #candidates = json.load(f)
    for line in f:
        line = line.strip()

        if line:  # Skip empty lines
            candidates.append(json.loads(line))

parser = CandidateParser()

parsed_candidates = []

for candidate in candidates:

    parsed = parser.parse(candidate)

    parsed_candidates.append(parsed)

# Save parsed candidates
with open("D:\Resume Parser and Ranker - Hackathon\Project\data\parsed_candidates.json", "w", encoding="utf-8") as f:

    json.dump(parsed_candidates, f, indent=4, cls=NumpyEncoder)

print("Done!")


normalizer = CandidateNormalizer()

normalized_candidates = []

for candidate in parsed_candidates:

    normalized_candidates.append(
        normalizer.normalize(candidate)
    )


skill_matcher = SkillMatcher()
experience_matcher = ExperienceMatcher()
embedding_matcher = EmbeddingMatcher()


scorer = Scorer()

scored_candidates = []

for candidate in normalized_candidates:

    features = {

        "skills":
            skill_matcher.match(jd, candidate),

        "experience":
            experience_matcher.match(jd, candidate),

        "embedding":
            embedding_matcher.match(jd, candidate),

        # Later
        "behaviour": 1,
        "company": 1,
        "location": 1
    }

    score = scorer.score(features)

    candidate["score"] = score
    candidate["features"] = features

    scored_candidates.append(candidate)




ranker = Ranker()

ranked_candidates = ranker.rank(scored_candidates)

from submission import save_submission

save_submission(
    jd,
    ranked_candidates,
    "data/submission.csv"
)



with open(
    "data/ranked_candidates.json",
    "w",
    encoding="utf-8"
) as f:

    json.dump(
        ranked_candidates,
        f,
        indent=4,
        cls=NumpyEncoder
    )

top_100 = ranked_candidates[:100]

with open(
    "data/top_100.json",
    "w",
    encoding="utf-8"
) as f:

    json.dump(
        top_100,
        f,
        indent=4,
        cls=NumpyEncoder
    )





"""




import os
import sys
import json
import numpy as np
from pprint import pprint
from docx import Document

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from resume_parser.jd_parser import JDParser
from resume_parser.candidate_parser import CandidateParser

from normalizer.candidate_normalizer import CandidateNormalizer

from matcher.skill_matcher import SkillMatcher
from matcher.experience_matcher import ExperienceMatcher
from matcher.embedding_matcher import EmbeddingMatcher

from scorer.scorer import Scorer
from ranker.ranker import Ranker
from submission import save_submission


# -----------------------------
# JSON Encoder
# -----------------------------

class NumpyEncoder(json.JSONEncoder):
    def default(self, obj):

        if isinstance(obj, np.ndarray):
            return obj.tolist()

        if isinstance(obj, (np.int32, np.int64)):
            return int(obj)

        if isinstance(obj, (np.float32, np.float64)):
            return float(obj)

        return super().default(obj)


# -----------------------------
# Read Job Description
# -----------------------------

doc = Document(r"D:\Resume Parser and Ranker - Hackathon\Project\data\job_description.docx")

text = "\n".join(
    paragraph.text
    for paragraph in doc.paragraphs
)

jd_parser = JDParser()

jd = jd_parser.parse(text)

print("\nParsed JD")
pprint(jd)


# -----------------------------
# Initialize Components
# -----------------------------

candidate_parser = CandidateParser()

normalizer = CandidateNormalizer()

skill_matcher = SkillMatcher()

experience_matcher = ExperienceMatcher()

embedding_matcher = EmbeddingMatcher()

scorer = Scorer()

ranker = Ranker()


# -----------------------------
# Process Candidates
# -----------------------------

scored_candidates = []

print("\nProcessing candidates...\n")

with open(
    r"D:\Resume Parser and Ranker - Hackathon\Project\data\candidates.jsonl",
    "r",
    encoding="utf-8"
) as f:

    for index, line in enumerate(f, start=1):

        line = line.strip()

        if not line:
            continue

        candidate = json.loads(line)

        # ---------------- Parse ----------------

        candidate = candidate_parser.parse(candidate)

        # ---------------- Normalize ----------------

        candidate = normalizer.normalize(candidate)

        # ---------------- Matching ----------------

        features = {

            "skills":
                skill_matcher.match(jd, candidate),

            "experience":
                experience_matcher.match(jd, candidate),

            "embedding":
                embedding_matcher.match(jd, candidate),

            "behaviour": 1,

            "company": 1,

            "location": 1
        }

        score = scorer.score(features)

        candidate["features"] = features

        candidate["score"] = score

        # Embedding no longer required
        candidate.pop("embedding", None)

        scored_candidates.append(candidate)

        if index % 1000 == 0:
            print(f"Processed {index} candidates")


print("\nFinished Processing.\n")


# -----------------------------
# Ranking
# -----------------------------

ranked_candidates = ranker.rank(scored_candidates)


# -----------------------------
# Save Ranked Candidates
# -----------------------------

with open(
    r"D:\Resume Parser and Ranker - Hackathon\Project\data\ranked_candidates.json",
    "w",
    encoding="utf-8"
) as f:

    json.dump(
        ranked_candidates,
        f,
        cls=NumpyEncoder,
        separators=(",", ":")
    )


# -----------------------------
# Save Top 100
# -----------------------------

top_100 = ranked_candidates[:100]

with open(
    r"D:\Resume Parser and Ranker - Hackathon\Project\data\top_100.json",
    "w",
    encoding="utf-8"
) as f:

    json.dump(
        top_100,
        f,
        cls=NumpyEncoder,
        separators=(",", ":")
    )


# -----------------------------
# Save Submission CSV
# -----------------------------

save_submission(
    jd,
    top_100,
    r"D:\Resume Parser and Ranker - Hackathon\Project\data\submission.csv"
)


print("\nTop Candidate\n")
pprint(top_100[0])

print("\nDone!")